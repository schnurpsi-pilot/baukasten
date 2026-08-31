# -*- coding: utf-8 -*-
"""Orchestrierung: aus einer Spezifikation entsteht das fertige ZIP.

Ablauf entspricht §13: Artefakte erzeugen, Checkliste §14.1 abarbeiten,
Historie fortschreiben (§11.7), Sammelausgabe als ein ZIP (§18.7).
"""
import datetime as dt
import json
import os
import shutil
import subprocess
import zipfile

from docx import Document
from docx.shared import Pt, Cm

from . import docxbau as B
from . import dokumente as DOK
from . import layout as L
from . import objekte as O
from . import pruefung as P
from . import reihe as R
from . import xlsxbau as X


# ------------------------------------------------------- Aufgabendateien
def _bloecke_schreiben(doc, bloecke, fettbegriffe=()):
    """Baut ein Textverarbeitungsdokument aus einer Blockliste."""
    for block in bloecke:
        typ = block["typ"]
        if typ == "absatz":
            p = B.absatz(doc, "", nach=block.get("nach", 6),
                         vor=block.get("vor", 0),
                         blocksatz=block.get("blocksatz", False),
                         zentriert=block.get("zentriert", False))
            teile = block.get("teile") or [(block.get("text", ""), {})]
            for text, stil in teile:
                r = p.add_run(text)
                r.font.name = B.ARIAL
                r.font.size = Pt(stil.get("groesse", block.get("groesse", 11)))
                r.font.bold = stil.get("fett", block.get("fett", False))
        elif typ == "ueberschrift":
            B.ueberschrift(doc, block["text"], block.get("groesse", 12),
                           vor=block.get("vor", 12), nach=block.get("nach", 6),
                           zentriert=block.get("zentriert", False))
        elif typ == "aufzaehlung":
            for t in block["zeilen"]:
                p = doc.add_paragraph(style="List Bullet")
                r = p.add_run(t)
                r.font.name = B.ARIAL
                r.font.size = Pt(block.get("groesse", 11))
                p.paragraph_format.left_indent = Cm(block.get("einzug", 0.5))
                p.paragraph_format.first_line_indent = Cm(-block.get("einzug", 0.5))
                p.paragraph_format.space_after = Pt(block.get("nach", 3))
        elif typ == "tabelle":
            # Feste Breiten, wenn die Aufgabe sie messbar vorgibt (§6.2) —
            # sonst misst das Layout die Spalten selbst aus (§18.3).
            if block.get("breiten"):
                breiten = block["breiten"]
                groesse = block.get("pt", 10)
            else:
                breiten, groesse = L.auto_breiten(
                    block["kopf"], block["zeilen"],
                    gesamt_cm=block.get("gesamt", 16.0), pt=block.get("pt", 10))
            B.tabelle(doc, block["kopf"], block["zeilen"], breiten, groesse=groesse)
        elif typ == "wordart":
            p = B.absatz(doc, "", nach=block.get("nach", 12),
                         vor=block.get("vor", 0), zentriert=True)
            O.wordart(p, block["text"], breite_pt=block.get("breite_pt", 340),
                      hoehe_pt=block.get("hoehe_pt", 42),
                      groesse_pt=block.get("groesse_pt", 36))
        elif typ == "leer":
            B.absatz(doc, "", nach=block.get("nach", 0))
        else:
            raise KeyError(f"Unbekannter Dokumentblock: {typ}")
    return doc


def _dokument_bauen(definition, pfad, fussnoten=None, fusszeile=False,
                    autotext=None):
    doc = Document()
    B.grundformat(doc)
    if autotext:
        B.autotext_fusszeile(doc, felder=autotext)
    elif fusszeile:
        B.seitenzahl_fusszeile(doc)
    _bloecke_schreiben(doc, definition)
    doc.save(pfad)
    if fussnoten:
        B.fussnoten_einfuegen(pfad, fussnoten)
    return pfad


def _vorlagendatei_bauen(block, vorlagen_ordner, pfad):
    doc = B.vorlage_oeffnen(vorlagen_ordner, block["form"], pfad)
    if block["form"] == "brief":
        B.brief_fuellen(doc, **block["inhalt"])
    else:
        B.metadaten_vorlage_fuellen(doc, block["inhalt"].get("felder", {}),
                                    block["inhalt"].get("koerper"))
    doc.save(pfad)
    return pfad


# ------------------------------------------------------------------- Historie
def historie_fortschreiben(spec, alte_datei, zielpfad, plan=None,
                           planpfad=None):
    """§11.7: übergebene Historie einlesen, neuen Satz ergänzen, zurückschreiben.

    Fortgeschrieben wird nicht nur die Satzliste, sondern auch das
    Reihenwissen: Quelle des Plans, Ergebnis der Sperrprüfung (§10.1, §10.2)
    und die nach §10.3 noch offenen Pflichtelemente. Ohne Reihenplan bleiben
    die planabhängigen Felder unverändert stehen, statt zu verschwinden.
    """
    if alte_datei and os.path.exists(alte_datei):
        historie = json.load(open(alte_datei, encoding="utf-8"))
        eingelesen = True
    else:
        historie = {"reihe": spec["meta"].get("reihe", "AP1"), "saetze": []}
        eingelesen = False

    # Der Stand ist der Tag der Fortschreibung, nicht ein Datum aus der
    # Satzspezifikation — sonst altert die Historie mit der Vorlage mit.
    historie["stand"] = dt.date.today().isoformat()
    if planpfad:
        historie["quelle_reihenplan"] = os.path.basename(planpfad)
    if spec["meta"].get("hinweis"):
        historie["hinweis"] = spec["meta"]["hinweis"]

    eintrag = dict(spec["historie_eintrag"])
    eintrag["sperrlisten_check"] = R.sperrlisten_check(eintrag, historie, plan)
    spec["sperrlisten_check"] = eintrag["sperrlisten_check"]

    # Der eigene Eintrag wird ersetzt, der unter altem Namen geführte
    # Vorgänger entfernt — sonst stünde derselbe Satz zweimal in der Reihe.
    ersetzt = {spec["meta"]["satzname"], eintrag.get("vorgaenger_satzname")}
    ersetzt.discard(None)
    historie["saetze"] = [s for s in historie["saetze"]
                          if s.get("satzname") not in ersetzt]
    historie["saetze"].append(eintrag)

    offen = R.pflichtbelegung(plan, historie, eintrag)
    if offen is not None:
        historie["pflichtelemente_offen_laut_plan"] = offen
        spec["pflichtelemente_offen"] = offen

    json.dump(historie, open(zielpfad, "w", encoding="utf-8"),
              ensure_ascii=False, indent=2)
    return eingelesen


# ----------------------------------------------------------------- Hauptlauf
def satz_bauen(spec, arbeitsordner, ausgabeordner, vorlagen_ordner,
               bilder_ordner=None, alte_historie=None, pdf=True,
               reihenplan=None):
    """Erzeugt alle Artefakte, prüft sie und packt das ZIP.

    Rückgabe: (zip_pfad, befund). Bei rotem Befund entsteht trotzdem ein ZIP,
    damit sich der Fehler nachvollziehen lässt — der Bericht sagt, was offen
    ist. Nichts wird ausgegeben, wenn ein Artefakt fehlt (§18.7).
    """
    os.makedirs(arbeitsordner, exist_ok=True)
    os.makedirs(ausgabeordner, exist_ok=True)
    name = spec["meta"]["satzname"]
    p = lambda datei: os.path.join(arbeitsordner, datei)

    erzeugt = []
    DOK.aufgabenbogen(spec, p(f"{name}_Aufgabenbogen.docx"))
    erzeugt.append(f"{name}_Aufgabenbogen.docx")
    DOK.materialheft(spec, p(f"{name}_Materialheft.docx"),
                     bilder_ordner=bilder_ordner)
    erzeugt.append(f"{name}_Materialheft.docx")

    teilnehmer, loesungen = [], []
    loesungspfad = teilnehmerpfad = None
    for datei in spec["dateien"]:
        praefix = datei["praefix"]
        if datei["art"] == "xlsx":
            tp = f"{name}_{praefix}_Teilnehmer.xlsx"
            lp = f"{name}_{praefix}_Loesung.xlsx"
            X.mappe_bauen(spec, False, p(tp))
            X.mappe_bauen(spec, True, p(lp))
            teilnehmerpfad, loesungspfad = p(tp), p(lp)
        elif datei["art"] == "dokument":
            tp = f"{name}_{praefix}_Teilnehmer.docx"
            lp = f"{name}_{praefix}_Loesung.docx"
            _dokument_bauen(datei["teilnehmer"], p(tp))
            _dokument_bauen(datei["loesung"], p(lp),
                            fussnoten=datei.get("fussnoten"),
                            fusszeile=datei.get("fusszeile", False))
        elif datei["art"] == "dotx":
            # Die Teilnehmerdatei ist ein gewöhnliches Dokument — das
            # Speichern als Dokumentvorlage ist selbst Prüfungsleistung
            # (Anhang D.2). Die Lösung ist deshalb die fertige .dotx.
            tp = f"{name}_{praefix}_Teilnehmer.docx"
            lp = f"{name}_{praefix}_Loesung.dotx"
            _dokument_bauen(datei["teilnehmer"], p(tp))
            zwischen = p(f"{name}_{praefix}_Loesung_roh.docx")
            _dokument_bauen(datei["loesung"], zwischen,
                            fussnoten=datei.get("fussnoten"),
                            fusszeile=datei.get("fusszeile", False),
                            autotext=datei.get("autotext"))
            O.als_dotx(zwischen, p(lp))
            os.remove(zwischen)
        elif datei["art"] == "vorlage":
            tp = f"{name}_{praefix}_Teilnehmer.docx"
            lp = f"{name}_{praefix}_Loesung.docx"
            _vorlagendatei_bauen({"form": datei["form"],
                                  "inhalt": datei["teilnehmer"]},
                                 vorlagen_ordner, p(tp))
            _vorlagendatei_bauen({"form": datei["form"], "inhalt": datei["loesung"]},
                                 vorlagen_ordner, p(lp))
        else:
            raise KeyError(f"Unbekannte Dateiart: {datei['art']}")
        teilnehmer.append(tp)
        loesungen.append(lp)

    DOK.bewertungsbogen(spec, p(f"{name}_Bewertungsbogen.docx"))
    DOK.handreichung(spec, p(f"{name}_Handreichung_Lehrkraft.docx"))

    if pdf and shutil.which("soffice"):
        subprocess.run(["soffice", "--headless", "--convert-to", "pdf",
                        "--outdir", arbeitsordner, p(f"{name}_Materialheft.docx")],
                       check=True, capture_output=True, timeout=300)

    plan = R.plan_laden(reihenplan)
    spec["reihenplan_eingelesen"] = plan is not None
    hist_datei = f"{name}_Historie.json"
    eingelesen = historie_fortschreiben(spec, alte_historie, p(hist_datei),
                                        plan=plan, planpfad=reihenplan)
    spec["historie_eingelesen"] = eingelesen

    # Reihenfolge im Archiv nach §18.7.
    dateiliste = ([f"{name}_Aufgabenbogen.docx", f"{name}_Materialheft.docx"]
                  + ([f"{name}_Materialheft.pdf"] if pdf else [])
                  + teilnehmer + loesungen
                  + [f"{name}_Bewertungsbogen.docx",
                     f"{name}_Handreichung_Lehrkraft.docx", hist_datei])

    befund = P.alles_pruefen(spec, arbeitsordner, dateiliste,
                             loesungspfad=loesungspfad,
                             teilnehmerpfad=teilnehmerpfad,
                             erwartet=spec.get("erwartete_werte"))

    fehlend = [d for d in dateiliste if not os.path.exists(p(d))]
    if fehlend:
        raise RuntimeError(f"Kein ZIP: Artefakte fehlen — {fehlend} (§18.7)")

    zip_pfad = os.path.join(ausgabeordner, f"{name}.zip")
    with zipfile.ZipFile(zip_pfad, "w", zipfile.ZIP_DEFLATED) as z:
        for datei in dateiliste:
            z.write(p(datei), datei)
    return zip_pfad, befund
