# -*- coding: utf-8 -*-
"""Erzeugt die vier Textartefakte des Satzes.

Aufgabenbogen (§5.2, §18.4), Materialheft (§11.4), Bewertungsbogen (§9.3,
§18.3) und Handreichung (§11.5).
"""
import re

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.section import WD_ORIENT

from . import docxbau as B
from . import layout as L

TAB_RECHTS_CM = 15.8        # Textbreite abzüglich 0,2 cm (§18.3)

KENNZEICHNUNG = ("Prüfungssimulation, erstellt von WBS Training SE. Alle Aufgaben, "
                 "Sachverhalte und Lösungen sind eigenständig erstellt.")

HILFSMITTEL = ("Erlaubte Hilfsmittel: keine außer den bereitgestellten Dateien und "
               "Anlagen. Kein Taschenrechner, keine Formelsammlung.")

# Diese drei Sätze stehen wörtlich in den Allgemeinen Hinweisen (§5.2).
SATZ_FORMATPUNKTE = ("In jeder Aufgabe entfallen zusätzlich 2 Punkte auf die "
                     "Einhaltung der Formatvorgaben.")
SATZ_KOPIERFAEHIG = ("Verwenden Sie überall dort Formeln, die Sie nach unten kopieren "
                     "können, wo mehrere Zeilen gleichartig berechnet werden.")
SATZ_UMSATZSTEUER = ("Alle Beträge in den Anlagen und Dateien sind Nettobeträge; der "
                     "Umsatzsteuersatz beträgt 19 Prozent.")


def _runs_mit_auszeichnung(par, text, fettbegriffe, groesse=11):
    """Zeichnet Datei- und Blattnamen sowie @@...@@ fett aus (§18.3)."""
    if fettbegriffe:
        # Wortgrenzen sind Pflicht, sonst zeichnet ein kurzer Blattname jedes
        # Kompositum an: aus "Postarten" würde "Postarten" mit fettem Kopf.
        muster = "|".join(re.escape(b) for b in sorted(fettbegriffe, key=len,
                                                       reverse=True))
        teile = re.split(rf"(@@[^@]+@@|(?<![\wÄÖÜäöüß])(?:{muster})"
                         rf"(?![\wÄÖÜäöüß]))", text)
    else:
        teile = re.split(r"(@@[^@]+@@)", text)
    for teil in teile:
        if not teil:
            continue
        fett = False
        if teil.startswith("@@") and teil.endswith("@@"):
            teil, fett = teil[2:-2], True
        elif fettbegriffe and teil in fettbegriffe:
            fett = True
        r = par.add_run(teil)
        r.font.name = B.ARIAL
        r.font.size = Pt(groesse)
        r.font.bold = fett
    return par


# ------------------------------------------------------------- Aufgabenbogen
def aufgabenbogen(spec, pfad):
    m = spec["meta"]
    fettbegriffe = set(spec.get("fettbegriffe", []))
    doc = Document()
    B.grundformat(doc)
    B.seitenzahl_fusszeile(doc)

    # Seite 1: Deckblatt allein (§18.4). Nur die Felder aus §5.2 Punkt 1.
    B.absatz(doc, "", nach=0).paragraph_format.space_before = Pt(120)
    B.ueberschrift(doc, "Prüfungssimulation", 18, nach=6, zentriert=True, halte=False)
    for zeile in m["bezeichnung"]:
        B.ueberschrift(doc, zeile, 14, nach=2, zentriert=True, halte=False)
    doc.paragraphs[-1].paragraph_format.space_after = Pt(36)

    tab = B.feldtabelle(doc, [
        ("Satz", m["satzname"]),
        ("Bearbeitungszeit", f"{m['bearbeitungszeit']} Minuten"),
        ("Gesamtpunkte", f"{m['gesamtpunkte']} Punkte"),
        ("Aufgaben", f"{len(spec['aufgaben'])} Aufgaben"),
    ], breiten_cm=(6.0, 6.0))
    tab.alignment = WD_TABLE_ALIGNMENT.CENTER

    B.absatz(doc, "", nach=24)
    B.absatz(doc, HILFSMITTEL, nach=36, zentriert=True)
    B.absatz(doc, KENNZEICHNUNG, nach=0, groesse=10, kursiv=True, zentriert=True)

    # Seite 2: Vorspann in fester Reihenfolge (§18.4).
    B.seitenumbruch(doc)
    B.ueberschrift(doc, "Allgemeine Hinweise", 14, nach=8)
    for titel, text in spec["hinweise"]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(0.7)
        p.paragraph_format.first_line_indent = Cm(-0.7)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(titel + ": ")
        r.font.name = B.ARIAL; r.font.size = Pt(11); r.font.bold = True
        r = p.add_run(text)
        r.font.name = B.ARIAL; r.font.size = Pt(11)

    B.ueberschrift(doc, "Formatvorgaben für die Bearbeitung", 14, vor=12, nach=4)
    p = B.absatz(doc, "Diese Vorgaben gelten für den gesamten Prüfungssatz. Die "
                      "einzelnen Aufgaben wiederholen sie nicht.", nach=4)
    p.paragraph_format.keep_with_next = True
    tab = B.tabelle(doc, ["Art der Angabe", "Format"], spec["formatvorgaben"],
                    [5.4, 10.6], groesse=10)
    for row in tab.rows:
        for cell in row.cells:
            for par in cell.paragraphs:
                if par.alignment == WD_ALIGN_PARAGRAPH.CENTER:
                    par.alignment = WD_ALIGN_PARAGRAPH.LEFT

    B.ueberschrift(doc, "Sachverhalt", 14, vor=12, nach=4)
    B.absatz(doc, spec["sachverhalt"], nach=0, blocksatz=True)

    # Jede Aufgabe auf neuer Seite (§18.4).
    for a in spec["aufgaben"]:
        B.seitenumbruch(doc)
        B.ueberschrift(doc, f"Aufgabe {a['nr']} – {a['titel']} ({a['punkte']} Punkte)",
                       14, nach=8)
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(10)
        _runs_mit_auszeichnung(p, a["einleitung"], fettbegriffe)

        for nr, text, punkte in a["teilaufgaben"]:
            p = doc.add_paragraph()
            pf = p.paragraph_format
            pf.left_indent = Cm(0.8)
            pf.first_line_indent = Cm(-0.8)
            pf.space_after = Pt(8)
            pf.keep_together = True
            pf.tab_stops.add_tab_stop(Cm(TAB_RECHTS_CM), WD_TAB_ALIGNMENT.RIGHT)
            r = p.add_run(f"{nr})\t")
            r.font.name = B.ARIAL; r.font.size = Pt(11)
            _runs_mit_auszeichnung(p, text, fettbegriffe)
            punktetext = f"({punkte} Punkte)"
            r = p.add_run()
            r.font.name = B.ARIAL; r.font.size = Pt(11)
            if L.punkte_umbrechen(text, punktetext):
                r.add_break()
            r = p.add_run(f"\t{punktetext}")
            r.font.name = B.ARIAL; r.font.size = Pt(11)

    doc.save(pfad)
    return pfad


# -------------------------------------------------------------- Materialheft
def materialheft(spec, pfad, bilder_ordner=None):
    m = spec["meta"]
    fettbegriffe = set(spec.get("fettbegriffe", []))
    doc = Document()
    B.grundformat(doc)
    B.seitenzahl_fusszeile(doc)
    doc.sections[0].header.paragraphs[0].text = ""

    B.ueberschrift(doc, "Materialheft", 18, nach=6, halte=False)
    B.ueberschrift(doc, m["satzname"], 14, nach=18, halte=False)
    B.absatz(doc, "Die Anlagen gehören zum Aufgabenbogen. Jede Aufgabe nennt die "
                  "Anlagen, die für ihre Bearbeitung gebraucht werden.",
             nach=12, blocksatz=True)
    B.tabelle(doc, ["Anlage", "Inhalt", "gehört zu"],
              [[f"Anlage {a['nr']}", a["titel"], a.get("gehoert_zu", "")]
               for a in spec["anlagen"]],
              [3.0, 8.6, 4.4], groesse=10)

    for anlage in spec["anlagen"]:
        quer = anlage.get("quer", False)
        sec = B.abschnitt(doc, quer=quer)
        titelzeile = f"Anlage {anlage['nr']} – {anlage['titel']}"
        B.kopfzeile(sec, titelzeile)
        B.ueberschrift(doc, titelzeile, 14, nach=10)
        gesamt = 24.7 if quer else 16.0

        for block in anlage["bloecke"]:
            typ = block["typ"]
            if typ == "text":
                B.absatz(doc, block["text"], nach=block.get("nach", 8),
                         blocksatz=True)
            elif typ == "ueberschrift":
                B.ueberschrift(doc, block["text"], block.get("groesse", 12),
                               vor=block.get("vor", 14), nach=6)
            elif typ == "felder":
                B.feldtabelle(doc, block["paare"],
                              breiten_cm=block.get("breiten", (3.4, gesamt - 3.4)))
                B.absatz(doc, "", nach=6)
            elif typ == "absaetze":
                for zeile in block["zeilen"]:
                    B.absatz(doc, zeile, nach=0 if zeile == "" else 6)
            elif typ == "tabelle":
                breiten, groesse = L.auto_breiten(
                    block["kopf"], block["zeilen"], gesamt_cm=gesamt,
                    pt=block.get("pt", 10))
                B.tabelle(doc, block["kopf"], block["zeilen"], breiten,
                          groesse=groesse,
                          zahlenspalten=tuple(block.get("zahlenspalten", ())))
            elif typ == "liste":
                B.aufzaehlung(doc, block["zeilen"],
                              runs_fkt=lambda p, t: _runs_mit_auszeichnung(
                                  p, t, fettbegriffe))
            elif typ == "quellen":
                for marke, text in block["eintraege"]:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Cm(1.0)
                    p.paragraph_format.first_line_indent = Cm(-1.0)
                    p.paragraph_format.space_after = Pt(6)
                    r = p.add_run(f"{marke}\t")
                    r.font.name = B.ARIAL; r.font.size = Pt(11); r.font.bold = True
                    r = p.add_run(text)
                    r.font.name = B.ARIAL; r.font.size = Pt(11)
            elif typ == "bild":
                pfad_bild = block["pfad"]
                if bilder_ordner and not pfad_bild.startswith("/"):
                    import os
                    pfad_bild = os.path.join(bilder_ordner, pfad_bild)
                doc.add_picture(pfad_bild, width=Cm(block.get("breite",
                                                              gesamt - 2.0)))
                doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                raise KeyError(f"Unbekannter Anlagenblock: {typ}")

    doc.save(pfad)
    return pfad


# ------------------------------------------------------------ Bewertungsbogen
BEWERTUNGSKOPF = ["Nr.", "Erwartete Leistung", "Lösungshinweis", "Punkte", "Toleranz"]
BEWERTUNGSBREITEN = [2.0, 6.2, 7.6, 1.7, 7.2]      # §18.3

VORSPANN_BEWERTUNG = (
    "Folgefehler werden nicht doppelt bestraft: Ein falsches Zwischenergebnis, das "
    "folgerichtig weiterverwendet wird, kostet nur an der Fehlerstelle Punkte. Die "
    "Funktionswahl ist frei; jede Funktion der Befehlsübersicht, die zum richtigen "
    "und kopierfähigen Ergebnis führt, zählt voll. Selbstaktualisierende Werte "
    "werden über die Formel bewertet, nicht über den angezeigten Wert."
)


def bewertungsbogen(spec, pfad):
    m = spec["meta"]
    doc = Document()
    B.grundformat(doc)
    sec = doc.sections[0]
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.page_width, sec.page_height = Cm(29.7), Cm(21.0)
    sec.left_margin = sec.right_margin = Cm(2.5)
    sec.top_margin = sec.bottom_margin = Cm(2.5)
    B.seitenzahl_fusszeile(doc)

    B.ueberschrift(doc, f"Bewertungsbogen {m['satzname']}", 14, nach=6)
    verteilung = " · ".join(f"Aufgabe {a['nr']}: {a['punkte']} Punkte"
                            for a in spec["aufgaben"])
    B.absatz(doc, f"Gesamtpunkte {m['gesamtpunkte']} · Bearbeitungszeit "
                  f"{m['bearbeitungszeit']} Minuten · {verteilung}", nach=4)
    B.absatz(doc, VORSPANN_BEWERTUNG, nach=10, blocksatz=True)

    zeilen = [[b["nr"], b["leistung"], b["hinweis"], b["punkte"], b["toleranz"]]
              for b in spec["bewertung"]]
    tab = B.tabelle(doc, BEWERTUNGSKOPF, zeilen, BEWERTUNGSBREITEN, groesse=8,
                    zahlenspalten=(3,))
    for ri, row in enumerate(tab.rows):
        if ri == 0:
            continue
        for ci, cell in enumerate(row.cells):
            for par in cell.paragraphs:
                if ci in (0, 1, 2, 4):
                    par.alignment = WD_ALIGN_PARAGRAPH.LEFT
                if ci == 0:
                    for r in par.runs:
                        r.font.bold = True

    B.absatz(doc, "", nach=6)
    B.absatz(doc, f"Summe: {sum(b['punkte'] for b in spec['bewertung'])} Punkte",
             nach=0, fett=True)
    doc.save(pfad)
    return pfad


# --------------------------------------------------------------- Handreichung
def handreichung(spec, pfad):
    m = spec["meta"]
    h = spec["handreichung"]
    doc = Document()
    B.grundformat(doc)
    B.seitenzahl_fusszeile(doc)

    B.ueberschrift(doc, f"Handreichung für die Lehrkraft – {m['satzname']}", 14,
                   nach=10)

    B.ueberschrift(doc, "1 Übersicht", 12, vor=8, nach=6)
    B.feldtabelle(doc, h["uebersicht"])

    B.ueberschrift(doc, "2 Zeitraster für die Besprechung", 12, vor=14, nach=6)
    B.absatz(doc, "Planungsrichtwert 60 Minuten. Je nach Gruppe kann die Besprechung "
                  "kürzer oder länger ausfallen.", nach=6, blocksatz=True)
    B.tabelle(doc, ["Abschnitt", "Minuten"],
              [[a, str(mi)] for a, mi in h["zeitraster"]],
              [13.0, 3.0], zahlenspalten=(1,))

    B.ueberschrift(doc, "3 Erwartete Stolperstellen", 12, vor=14, nach=6)
    for nr, text in h["stolperstellen"]:
        p = doc.add_paragraph()
        p.paragraph_format.left_indent = Cm(1.2)
        p.paragraph_format.first_line_indent = Cm(-1.2)
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(f"{nr}\t")
        r.font.name = B.ARIAL; r.font.size = Pt(11); r.font.bold = True
        r = p.add_run(text)
        r.font.name = B.ARIAL; r.font.size = Pt(11)

    B.ueberschrift(doc, "4 Typische Falschlösungen und ihre Bewertung", 12, vor=14,
                   nach=6)
    B.absatz(doc, "Die Abzüge stehen jeweils in der Spalte Toleranz des "
                  "Bewertungsbogens. Grundsatz: Folgefehler kosten nur an der "
                  "Fehlerstelle Punkte.", nach=6, blocksatz=True)
    B.tabelle(doc, ["Falschlösung", "Teilaufgabe", "Bewertung"],
              h["falschloesungen"], [5.6, 2.6, 7.8], groesse=10)

    B.ueberschrift(doc, "5 Fachliche Klärung zu den Rechenwegen", 12, vor=14, nach=6)
    for begriff, text in h["fachklaerung"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(begriff + ": ")
        r.font.name = B.ARIAL; r.font.size = Pt(11); r.font.bold = True
        r = p.add_run(text)
        r.font.name = B.ARIAL; r.font.size = Pt(11)

    B.ueberschrift(doc, "6 Anschlussübungen", 12, vor=14, nach=6)
    for titel, text in h["anschlussuebungen"]:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        r = p.add_run(titel + ": ")
        r.font.name = B.ARIAL; r.font.size = Pt(11); r.font.bold = True
        r = p.add_run(text)
        r.font.name = B.ARIAL; r.font.size = Pt(11)

    if h.get("dateihinweise"):
        B.ueberschrift(doc, "7 Hinweise zu den Lösungsdateien", 12, vor=14, nach=6)
        for text in h["dateihinweise"]:
            B.absatz(doc, text, nach=6, blocksatz=True)

    doc.save(pfad)
    return pfad
