# -*- coding: utf-8 -*-
"""docx-Baukasten für die Prüfungssätze.

Enthält alles, was python-docx nicht mitbringt: echte Fußnoten, feste
Tabellenbreiten, Seitenzahlfelder und die regelkonforme Befüllung der
Goldberg-Vorlagen (§11.6).
"""
import copy
import os
import re
import shutil
import zipfile

from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from . import layout as L

ARIAL = "Arial"


# ------------------------------------------------------------------ Grundlagen
def grundformat(doc, raender_cm=2.5, zeilenabstand=1.15, groesse=11):
    """Setzt Hausschrift, Zeilenabstand und Seitenränder (§18.3)."""
    st = doc.styles["Normal"]
    st.font.name = ARIAL
    st.font.size = Pt(groesse)
    rpr = st.element.get_or_add_rPr()
    rf = rpr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rpr.append(rf)
    for a in ("w:ascii", "w:hAnsi", "w:cs", "w:eastAsia"):
        rf.set(qn(a), ARIAL)
    st.paragraph_format.line_spacing = zeilenabstand
    st.paragraph_format.space_after = Pt(0)
    for s in doc.sections:
        s.left_margin = s.right_margin = Cm(raender_cm)
        s.top_margin = s.bottom_margin = Cm(raender_cm)
    return doc


def _feld(par, instr):
    for typ in ("begin",):
        r = par.add_run()
        fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), typ); r._r.append(fc)
    r = par.add_run()
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    r._r.append(it)
    r = par.add_run()
    fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "separate"); r._r.append(fc)
    par.add_run("1")
    r = par.add_run()
    fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "end"); r._r.append(fc)


def seitenzahl_fusszeile(doc, groesse=9):
    """Fußzeile zentriert im Format 'Seite X von Y' (§18.3)."""
    for s in doc.sections:
        p = s.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.text = ""
        p.add_run("Seite ")
        _feld(p, " PAGE ")
        p.add_run(" von ")
        _feld(p, " NUMPAGES ")
        for r in p.runs:
            r.font.name = ARIAL
            r.font.size = Pt(groesse)


def autotext_fusszeile(doc, groesse=9, felder=("dateiname", "datum", "seiten")):
    """Fußzeile aus Autotext-Feldern (Anhang D.2).

    Autotext meint in der Befehlsübersicht die einfügbaren Felder Dateiname,
    Datum, Seitenanzahl und Uhrzeit. Sie werden als echte Feldfunktionen
    gesetzt, nicht als getippter Text — nur dann aktualisieren sie sich und
    nur dann ist die Prüfungsleistung nachweisbar.
    """
    ANWEISUNG = {"dateiname": " FILENAME \\* MERGEFORMAT ",
                 "datum": " DATE \\@ \"dd.MM.yyyy\" ",
                 "uhrzeit": " TIME \\@ \"HH:mm\" ",
                 "seiten": None}
    for s in doc.sections:
        p = s.footer.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.text = ""
        for i, feld in enumerate(felder):
            if i:
                p.add_run("   \u00b7   ")
            if feld == "seiten":
                p.add_run("Seite ")
                _feld(p, " PAGE ")
                p.add_run(" von ")
                _feld(p, " NUMPAGES ")
            elif feld in ANWEISUNG:
                _feld(p, ANWEISUNG[feld])
            else:
                raise KeyError(f"Unbekanntes Autotext-Feld: {feld}")
        for r in p.runs:
            r.font.name = ARIAL
            r.font.size = Pt(groesse)
    return doc


def ueberschrift(doc, text, groesse=14, vor=0, nach=12, zentriert=False, halte=True):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = ARIAL; r.font.size = Pt(groesse); r.font.bold = True
    p.paragraph_format.space_before = Pt(vor)
    p.paragraph_format.space_after = Pt(nach)
    if zentriert:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if halte:
        p.paragraph_format.keep_with_next = True
    return p


def absatz(doc, text="", nach=6, vor=0, fett=False, kursiv=False, blocksatz=False,
           zentriert=False, groesse=11, einzug=None, haengend=None):
    p = doc.add_paragraph()
    if text:
        r = p.add_run(text)
        r.font.name = ARIAL; r.font.size = Pt(groesse)
        r.font.bold = fett; r.font.italic = kursiv
    p.paragraph_format.space_after = Pt(nach)
    p.paragraph_format.space_before = Pt(vor)
    if blocksatz:
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if zentriert:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    if einzug is not None:
        p.paragraph_format.left_indent = Cm(einzug)
    if haengend is not None:
        p.paragraph_format.first_line_indent = Cm(-haengend)
    return p


def aufzaehlung(doc, texte, einzug=0.5, nach=4, groesse=11, runs_fkt=None):
    for t in texte:
        p = doc.add_paragraph(style="List Bullet")
        if runs_fkt:
            runs_fkt(p, t)
        else:
            r = p.add_run(t)
            r.font.name = ARIAL; r.font.size = Pt(groesse)
        p.paragraph_format.left_indent = Cm(einzug)
        p.paragraph_format.first_line_indent = Cm(-einzug)
        p.paragraph_format.space_after = Pt(nach)


def seitenumbruch(doc):
    p = doc.add_paragraph()
    p.add_run().add_break(WD_BREAK.PAGE)
    p.paragraph_format.space_after = Pt(0)
    return p


def abschnitt(doc, quer=False, erste=False, raender_cm=2.5):
    """Neuer Abschnitt mit eigener Ausrichtung und eigener Kopfzeile."""
    sec = doc.sections[-1] if erste else doc.add_section(WD_SECTION.NEW_PAGE)
    if quer:
        sec.orientation = WD_ORIENT.LANDSCAPE
        sec.page_width, sec.page_height = Cm(29.7), Cm(21.0)
    else:
        sec.orientation = WD_ORIENT.PORTRAIT
        sec.page_width, sec.page_height = Cm(21.0), Cm(29.7)
    sec.left_margin = sec.right_margin = Cm(raender_cm)
    sec.top_margin = sec.bottom_margin = Cm(raender_cm)
    return sec


def kopfzeile(sec, text, groesse=10, fett=True):
    sec.header.is_linked_to_previous = False
    p = sec.header.paragraphs[0]
    p.text = ""
    r = p.add_run(text)
    r.font.name = ARIAL; r.font.size = Pt(groesse); r.font.bold = fett
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    return p


# -------------------------------------------------------------------- Tabellen
# Spaltenüberschriften, die eine Kennung ankündigen. Mit einer Personal-,
# Artikel- oder Lieferantennummer wird nicht gerechnet — sie ist Text, auch
# wenn sie nur aus Ziffern besteht, und wird deshalb nie rechtsbündig
# gesetzt. Geprüft wird die Überschrift, nicht der Zellinhalt: Ob 4101 eine
# Nummer oder eine Menge ist, verrät allein der Spaltenkopf.
# Wortbestandteile: schlagen überall im Kopf an, auch zusammengeschrieben.
_KENNUNG_TEIL = ("nummer", "kennung", "kürzel", "iban", "zeichen")
# Eigenständige Wörter: nur an der Wortgrenze, sonst gäbe es Fehltreffer.
_KENNUNG_WORT = {"nr", "id", "code", "konto", "ean", "bln"}


def _ist_kennungsspalte(ueberschrift):
    """Trägt die Spalte laut Überschrift eine Kennung statt einer Rechengröße?

    Erkannt werden getrennte Formen ("Lief.-Nr.", "Personal-nummer") wie
    zusammengeschriebene ("Angebotsnummer", "Kundennummer"). Deshalb wird
    für die langen Wörter im Text gesucht, für die kurzen und mehrdeutigen
    ("nr", "id", "code") dagegen nur an der Wortgrenze — sonst schlüge
    "id" in "Liquidität" an.
    """
    if not ueberschrift:
        return False
    text = str(ueberschrift).lower().replace("\n", " ").replace("-", " ")
    if any(w in text for w in _KENNUNG_TEIL):
        return True
    woerter = [w.strip(".:") for w in text.split()]
    return any(w in _KENNUNG_WORT for w in woerter)


_RE_DATUM = re.compile(r"^\d{2}\.\d{2}\.\d{4}$")
# Zahl mit Tausenderpunkt (1.250,00) oder ohne Gliederung (2026, 120) —
# Jahreszahlen tragen laut §5.7 nie einen Tausenderpunkt.
_RE_ZAHL = re.compile(r"^-?(\d{1,3}(\.\d{3})+|\d+)(,\d+)?"
                      r"(\s?(EUR|%|Stück|Tage|Min|Std|kg|km))?$")


def _ist_zahl_oder_datum(wert):
    """Rechtsbündig gehört, was gerechnet oder datiert ist (§5.7, §18.3).

    Erkannt werden Datumsangaben im Format TT.MM.JJJJ, reine Zahlen sowie
    Zahlen mit nachgestellter Einheit oder Währung. Alles andere — auch
    Angebotsnummern wie AN-2026-4417 — bleibt Text und wird nicht
    ausgerichtet, obwohl es Ziffern enthält.
    """
    w = str(wert).strip()
    if not w:
        return False
    return bool(_RE_DATUM.match(w) or _RE_ZAHL.match(w))


def tabelle(doc, kopf, zeilen, breiten_cm, groesse=10, mit_kopf=True,
            zahlenspalten=(), erste_spalte_links=True):
    """Tabelle mit festen Spaltenbreiten (§18.3).

    Feste Breiten brauchen beides: tblLayout=fixed plus Breite an jeder
    Zelle UND jeder Spalte. Ohne das verteilt Word/LibreOffice neu und
    Spaltenüberschriften brechen mitten im Wort.
    """
    zeilen = list(zeilen)
    tab = doc.add_table(rows=(1 if mit_kopf else 0) + len(zeilen),
                        cols=len(breiten_cm))
    if mit_kopf:
        for i, t in enumerate(kopf):
            tab.rows[0].cells[i].text = str(t)
    versatz = 1 if mit_kopf else 0
    for ri, row in enumerate(zeilen):
        for ci, v in enumerate(row):
            tab.rows[ri + versatz].cells[ci].text = "" if v is None else str(v)

    # Zahlen- und Datumsspalten werden erkannt, nicht aufgezählt (§5.7):
    # Eine Spalte, deren gefüllte Datenzellen durchgehend Zahlen, Beträge
    # oder Datumsangaben tragen, steht rechtsbündig. Ohne die Erkennung
    # müsste jede Satzspezifikation ihre Datumsspalten einzeln benennen —
    # und eine vergessene Spalte fiele stillschweigend auf Zentriert zurück.
    zahlenspalten = set(zahlenspalten or ())
    kennungsspalten = set()
    for ci in range(len(breiten_cm)):
        if mit_kopf and ci < len(kopf) and _ist_kennungsspalte(kopf[ci]):
            zahlenspalten.discard(ci)
            kennungsspalten.add(ci)
            continue
        werte = [str(row[ci]).strip() for row in zeilen
                 if ci < len(row) and row[ci] not in (None, "")]
        if werte and all(_ist_zahl_oder_datum(w) for w in werte):
            zahlenspalten.add(ci)

    tab.style = "Table Grid"
    tab.alignment = WD_TABLE_ALIGNMENT.LEFT
    tab.autofit = False
    tblPr = tab._tbl.tblPr
    for alt in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(alt)
    lay = OxmlElement("w:tblLayout"); lay.set(qn("w:type"), "fixed")
    tblPr.append(lay)
    for i, b in enumerate(breiten_cm):
        tab.columns[i].width = Cm(b)

    for ri, row in enumerate(tab.rows):
        row._tr.get_or_add_trPr().append(OxmlElement("w:cantSplit"))
        kopfzeile_ist = (ri == 0 and mit_kopf)
        if kopfzeile_ist:
            row._tr.get_or_add_trPr().append(OxmlElement("w:tblHeader"))
        for ci, cell in enumerate(row.cells):
            cell.width = Cm(breiten_cm[ci])
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if kopfzeile_ist:
                tcPr = cell._tc.get_or_add_tcPr()
                shd = OxmlElement("w:shd")
                shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), "D9D9D9")
                tcPr.append(shd)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(1)
                p.paragraph_format.space_before = Pt(1)
                p.paragraph_format.line_spacing = 1.0
                for r in p.runs:
                    r.font.name = ARIAL
                    r.font.size = Pt(groesse)
                    r.font.bold = kopfzeile_ist
                if kopfzeile_ist:
                    p.alignment = (WD_ALIGN_PARAGRAPH.LEFT if ci == 0
                                   else WD_ALIGN_PARAGRAPH.CENTER)
                elif ci in zahlenspalten:
                    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
                elif ci == 0 and erste_spalte_links:
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                elif ci in kennungsspalten:
                    # Kennungen stehen links, gleich in welcher Spalte:
                    # untereinander bleiben sie so vergleichbar.
                    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
                else:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    return tab


def feldtabelle(doc, paare, breiten_cm=(4.6, 11.4), groesse=11):
    """Zweispaltige Tabelle Bezeichnung/Wert, linke Spalte fett."""
    tab = tabelle(doc, [], [[a, b] for a, b in paare], list(breiten_cm),
                  groesse=groesse, mit_kopf=False)
    for row in tab.rows:
        for c, cell in enumerate(row.cells):
            for par in cell.paragraphs:
                par.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for r in par.runs:
                    r.font.bold = (c == 0)
    return tab


# ------------------------------------------------------------------- Fußnoten
_FN_KOPF = (
    '<?xml version="1.0" encoding="UTF-8" standalone="yes"?>\n'
    '<w:footnotes xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main">'
    '<w:footnote w:type="separator" w:id="-1"><w:p><w:pPr><w:spacing w:after="0" '
    'w:line="240" w:lineRule="auto"/></w:pPr><w:r><w:separator/></w:r></w:p></w:footnote>'
    '<w:footnote w:type="continuationSeparator" w:id="0"><w:p><w:pPr><w:spacing '
    'w:after="0" w:line="240" w:lineRule="auto"/></w:pPr><w:r><w:continuationSeparator/>'
    '</w:r></w:p></w:footnote>'
)


def _escape(t):
    return t.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _fn_koerper(fid, text):
    return (
        f'<w:footnote w:id="{fid}"><w:p><w:pPr><w:spacing w:after="0" w:line="240" '
        f'w:lineRule="auto"/></w:pPr>'
        f'<w:r><w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>'
        f'<w:vertAlign w:val="superscript"/></w:rPr><w:footnoteRef/></w:r>'
        f'<w:r><w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial"/><w:sz w:val="18"/>'
        f'</w:rPr><w:t xml:space="preserve"> {_escape(text)}</w:t></w:r></w:p></w:footnote>'
    )


def fussnoten_einfuegen(docx_pfad, marker_texte):
    """Ersetzt Platzhalter-Runs durch echte Fußnoten.

    python-docx kann keine Fußnoten. Deshalb wird das fertige docx entpackt,
    footnotes.xml erzeugt und in Content-Types und Relationships eingehängt.
    marker_texte: {"@@FN1@@": "Quellentext", ...} in gewünschter Reihenfolge.
    Der Marker muss in einem eigenen Run stehen.
    """
    tmp = docx_pfad + ".unz"
    if os.path.isdir(tmp):
        shutil.rmtree(tmp)
    os.makedirs(tmp)
    with zipfile.ZipFile(docx_pfad) as z:
        z.extractall(tmp)

    doc_p = os.path.join(tmp, "word", "document.xml")
    doc_xml = open(doc_p, encoding="utf-8").read()

    fn_xml = _FN_KOPF
    for fid, (marker, text) in enumerate(marker_texte.items(), start=1):
        ref = ('<w:r><w:rPr><w:rFonts w:ascii="Arial" w:hAnsi="Arial"/>'
               f'<w:vertAlign w:val="superscript"/></w:rPr>'
               f'<w:footnoteReference w:id="{fid}"/></w:r>')
        pat = re.compile(r'<w:r>(?:(?!</w:r>).)*?' + re.escape(marker) + r'.*?</w:r>',
                         re.S)
        doc_xml, n = pat.subn(ref, doc_xml, count=1)
        if n != 1:
            raise RuntimeError(f"Fußnotenmarker {marker} nicht eindeutig gefunden "
                               f"(Treffer: {n}). Er muss in einem eigenen Run stehen.")
        fn_xml += _fn_koerper(fid, text)
    fn_xml += "</w:footnotes>"

    open(doc_p, "w", encoding="utf-8").write(doc_xml)
    open(os.path.join(tmp, "word", "footnotes.xml"), "w", encoding="utf-8").write(fn_xml)

    ct_p = os.path.join(tmp, "[Content_Types].xml")
    ct = open(ct_p, encoding="utf-8").read()
    if "footnotes+xml" not in ct:
        ct = ct.replace("</Types>",
                        '<Override PartName="/word/footnotes.xml" ContentType='
                        '"application/vnd.openxmlformats-officedocument.'
                        'wordprocessingml.footnotes+xml"/></Types>')
        open(ct_p, "w", encoding="utf-8").write(ct)

    rel_p = os.path.join(tmp, "word", "_rels", "document.xml.rels")
    rel = open(rel_p, encoding="utf-8").read()
    if "footnotes.xml" not in rel:
        ids = [int(m) for m in re.findall(r'Id="rId(\d+)"', rel)]
        neu = max(ids) + 1 if ids else 1
        rel = rel.replace("</Relationships>",
                          f'<Relationship Id="rId{neu}" Type="http://schemas.'
                          'openxmlformats.org/officeDocument/2006/relationships/footnotes"'
                          ' Target="footnotes.xml"/></Relationships>')
        open(rel_p, "w", encoding="utf-8").write(rel)

    os.remove(docx_pfad)
    with zipfile.ZipFile(docx_pfad, "w", zipfile.ZIP_DEFLATED) as z:
        for root, _d, files in os.walk(tmp):
            for f in files:
                voll = os.path.join(root, f)
                z.write(voll, os.path.relpath(voll, tmp))
    shutil.rmtree(tmp)


# ------------------------------------------------------- Goldberg-Wordvorlagen
VORLAGENDATEI = {
    "brief": "Brief_Goldberg_DIN5008.docx",
    "email": "E-Mail_Goldberg_DIN5008.docx",
    "aktenvermerk": "Aktenvermerk_Goldberg_DIN5008.docx",
    "einladung": "Einladung_Goldberg_DIN5008.docx",
    "protokoll": "Protokoll_Goldberg_DIN5008.docx",
    "rundschreiben": "Rundschreiben_Goldberg_DIN5008.docx",
    "mitteilung": "Innerbetriebliche_Mitteilung_Goldberg_DIN5008.docx",
}

# Zeilenindex der Beschriftungen in der Metadaten-Tabelle je Vorlage (§11.6).
BRIEF_FELDER = {
    "ihr_zeichen": 0, "ihre_nachricht_vom": 1, "unser_zeichen": 2,
    "unsere_nachricht_vom": 3, "name": 5, "telefon": 6, "telefax": 7,
    "email": 8, "datum": 10,
}


def _zelle_setzen(par, text, fett=False):
    """Schreibt Text in einen vorhandenen Absatz, ohne die Schrift zu setzen.

    Schriftart und -größe kommen aus der Formatvorlage der Vorlage; eigenes
    Setzen würde das Corporate Design überschreiben (§11.6).
    """
    for r in list(par.runs):
        r._r.getparent().remove(r._r)
    r = par.add_run(text)
    if fett:
        r.font.bold = True
    return par


def vorlage_oeffnen(vorlagen_ordner, form, zielpfad):
    """Kopiert die Goldberg-Vorlage und öffnet die Kopie. Nie nachbauen."""
    if form not in VORLAGENDATEI:
        raise KeyError(f"Unbekannte Kommunikationsform: {form}")
    quelle = os.path.join(vorlagen_ordner, VORLAGENDATEI[form])
    if not os.path.exists(quelle):
        raise FileNotFoundError(
            f"Vorlage fehlt: {quelle}. Nach §11.6 wird kein Ersatzlayout erzeugt — "
            "Vorlage anfordern.")
    shutil.copy(quelle, zielpfad)
    return Document(zielpfad)


# Nutzbare Breite der Wertespalte im Informationsblock des Briefes (§11.6).
# Die Spalte misst in der Goldberg-Vorlage 3,5 cm; abzüglich der beidseitigen
# Zellinnenabstände von je 0,19 cm bleiben rund 3,1 cm für den Text.
INFOBLOCK_BREITE_CM = 3.1
INFOBLOCK_MIN_PT = 6.0


def _einzeilig_einpassen(par, text, breite_cm, start_pt=11.0,
                         min_pt=INFOBLOCK_MIN_PT):
    """Verkleinert den Text stufenweise, bis er in eine Zeile passt.

    Betrifft vor allem lange E-Mail-Adressen im Informationsblock: Bricht
    die Adresse um, verschiebt sie alle folgenden Beschriftungen nach unten
    und das Datum steht nicht mehr auf seiner Zeile. Untergrenze ist 6 pt —
    darunter wird nicht weiter verkleinert, sondern der Umbruch in Kauf
    genommen, weil unlesbar schlechter ist als zweizeilig.

    Rückgabe: die gesetzte Schriftgröße in pt.
    """
    pt = start_pt
    while pt > min_pt and L.breite_cm(text, pt) > breite_cm:
        pt -= 0.5
    pt = max(pt, min_pt)
    for r in par.runs:
        r.font.size = Pt(pt)
    # Ein Umbruch innerhalb der Adresse bleibt sonst über die Silbentrennung
    # möglich; das Feld ist eine zusammengehörige Angabe.
    par.paragraph_format.keep_together = True
    return pt


def brief_fuellen(doc, empfaenger=None, infoblock=None, betreff=None,
                  absaetze=None, schluss=None, leerzeile_nach_betreff=True):
    """Befüllt die Briefvorlage: Anschriftfeld, Informationsblock, Fließtext.

    Der Brief ist der Sonderfall ohne Körpertabelle (§11.6): Anschriftfeld
    und Informationsblock liegen in einer 1x4-Tabelle, der Text steht in
    Fließtextabsätzen darunter. Es wird nur in vorhandene Zellen geschrieben.
    """
    zellen = doc.tables[0].rows[0].cells
    anschrift, _abstand, _label, werte = zellen

    if infoblock:
        for feld, wert in infoblock.items():
            idx = BRIEF_FELDER[feld]
            par = _zelle_setzen(werte.paragraphs[idx], wert)
            # Jeder Wert des Blocks bleibt einzeilig — geprüft wird die
            # Gattung, nicht das Feld E-Mail: Auch ein langer Name oder
            # eine gegliederte Telefonnummer darf die Zeilen nicht schieben.
            if L.breite_cm(wert, 11) > INFOBLOCK_BREITE_CM:
                _einzeilig_einpassen(par, wert, INFOBLOCK_BREITE_CM)

    if empfaenger:
        # Zeilen 0 bis 2 bleiben für Zusätze und Vermerke frei (DIN 5008:2020).
        for i, zeile in enumerate(empfaenger):
            _zelle_setzen(anschrift.paragraphs[3 + i], zeile)

    if betreff is None and absaetze is None:
        return doc

    ps = doc.paragraphs
    if betreff is not None:
        _zelle_setzen(ps[0], betreff, fett=True)
        _zelle_setzen(ps[1], "")
        if leerzeile_nach_betreff:
            ps[1]._p.addnext(copy.deepcopy(ps[1]._p))
    start = 3 if leerzeile_nach_betreff else 2
    if absaetze:
        _zelle_setzen(doc.paragraphs[start], absaetze[0])
        for zeile in absaetze[1:]:
            doc.add_paragraph(zeile)
    if schluss:
        p = doc.add_paragraph(schluss[0])
        p.paragraph_format.space_after = Pt(33)     # Unterschriftsraum §18.4
        for zeile in schluss[1:]:
            doc.add_paragraph(zeile)
    return doc


def metadaten_vorlage_fuellen(doc, felder, koerper=None):
    """Befüllt die übrigen Vorlagen (E-Mail, Protokoll, Notiz, ...).

    felder: {"Betreff": "…", "An": "…"} — die Beschriftung steuert die Zeile.
    koerper: Liste von Absätzen für die einzellige Körpertabelle.
    """
    tab = doc.tables[0]
    beschriftungen = {}
    for row in tab.rows:
        for ci, cell in enumerate(row.cells[:-1]):
            for pi, par in enumerate(cell.paragraphs):
                text = par.text.strip().rstrip(":")
                if text:
                    beschriftungen[text.lower()] = (row, ci + 1, pi)
    for name, wert in felder.items():
        treffer = beschriftungen.get(name.strip().rstrip(":").lower())
        if not treffer:
            raise KeyError(f"Feld '{name}' gibt es in dieser Vorlage nicht. "
                           f"Vorhanden: {sorted(beschriftungen)}")
        row, ci, pi = treffer
        _zelle_setzen(row.cells[ci].paragraphs[pi], wert)

    if koerper:
        koerpertab = doc.tables[-1]
        zelle = koerpertab.rows[0].cells[0]
        frei = zelle.paragraphs
        for i, zeile in enumerate(koerper):
            if i < len(frei):
                _zelle_setzen(frei[i], zeile)
            else:
                neu = copy.deepcopy(frei[-1]._p)
                frei[-1]._p.addnext(neu)
                frei = zelle.paragraphs
                _zelle_setzen(frei[i], zeile)
    return doc
